"""
Document extractor for various file formats (TXT, MD, PDF, DOCX, HTML).
"""
import logging
from pathlib import Path
from typing import Any, Dict, Union

logger = logging.getLogger(__name__)


def _extract_pdf(file_path: Path) -> Dict[str, Any]:
    """Extract text from PDF using available libraries with intelligent fallback."""
    text_pages = []
    page_count = 0
    errors = []

    # 1. Try pypdf (modern standard)
    try:
        import pypdf
        reader = pypdf.PdfReader(str(file_path))
        page_count = len(reader.pages)
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            text_pages.append({
                "section": f"Page {page_num + 1}",
                "text": page_text.strip()
            })
        full_text = "\n\n".join([p["text"] for p in text_pages if p["text"]])
        if full_text.strip():
            return {
                "text": full_text.strip(),
                "pages": text_pages,
                "page_count": page_count,
            }
    except Exception as e:
        errors.append(f"pypdf: {e}")

    # 2. Try PyPDF2
    try:
        import PyPDF2
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            text_pages = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                text_pages.append({
                    "section": f"Page {page_num + 1}",
                    "text": page_text.strip()
                })
            full_text = "\n\n".join([p["text"] for p in text_pages if p["text"]])
            if full_text.strip():
                return {
                    "text": full_text.strip(),
                    "pages": text_pages,
                    "page_count": page_count,
                }
    except Exception as e:
        errors.append(f"PyPDF2: {e}")

    # 3. Try pdfminer if available
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(str(file_path))
        if text.strip():
            return {
                "text": text.strip(),
                "pages": [{"section": "Full PDF Document", "text": text.strip()}],
                "page_count": 1,
            }
    except Exception as e:
        errors.append(f"pdfminer: {e}")

    # 4. Try fitz (PyMuPDF) if available
    try:
        import fitz
        doc = fitz.open(str(file_path))
        page_count = len(doc)
        text_pages = []
        for page_num, page in enumerate(doc):
            page_text = page.get_text() or ""
            text_pages.append({
                "section": f"Page {page_num + 1}",
                "text": page_text.strip()
            })
        full_text = "\n\n".join([p["text"] for p in text_pages if p["text"]])
        if full_text.strip():
            return {
                "text": full_text.strip(),
                "pages": text_pages,
                "page_count": page_count,
            }
    except Exception as e:
        errors.append(f"fitz: {e}")

    logger.warning(f"All PDF extractors failed for {file_path.name}: {'; '.join(errors)}")
    return {
        "text": "",
        "pages": [],
        "page_count": page_count,
        "error": f"PDF extraction failed: {'; '.join(errors)}"
    }


def _extract_docx(file_path: Path) -> Dict[str, Any]:
    """Extract text and tables from Word document (.docx)."""
    try:
        import docx
        doc = docx.Document(str(file_path))
        parts = []

        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                parts.append("\n".join(table_rows))

        full_text = "\n\n".join(parts)
        return {
            "text": full_text,
            "pages": [{"section": "Document Body", "text": full_text}]
        }
    except Exception as e:
        logger.warning(f"docx extractor error for {file_path}: {e}")
        return {
            "text": "",
            "pages": [],
            "error": str(e)
        }


def extract_text(file_path: Union[str, Path]) -> Dict[str, Any]:
    """
    Extract text from a document file safely.

    Args:
        file_path: Path to the document file

    Returns:
        Dictionary with extracted text, page sections, and metadata.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = file_path.suffix.lower()

    result = {
        "text": "",
        "metadata": {
            "file_name": file_path.name,
            "file_size": file_path.stat().st_size,
            "extension": extension,
        },
        "pages": [],
    }

    try:
        if extension == ".pdf":
            pdf_data = _extract_pdf(file_path)
            result["text"] = pdf_data.get("text", "")
            result["pages"] = pdf_data.get("pages", [])
            if "page_count" in pdf_data:
                result["metadata"]["page_count"] = pdf_data["page_count"]
            if "error" in pdf_data and not result["text"]:
                result["metadata"]["error"] = pdf_data["error"]

        elif extension in [".docx", ".doc"]:
            docx_data = _extract_docx(file_path)
            result["text"] = docx_data.get("text", "")
            result["pages"] = docx_data.get("pages", [])
            if "error" in docx_data and not result["text"]:
                result["metadata"]["error"] = docx_data["error"]

        elif extension in [".html", ".htm"]:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                raw_html = f.read()
            import re
            clean_text = re.sub(r"<script.*?</script>", " ", raw_html, flags=re.DOTALL | re.IGNORECASE)
            clean_text = re.sub(r"<style.*?</style>", " ", clean_text, flags=re.DOTALL | re.IGNORECASE)
            clean_text = re.sub(r"<[^>]+>", " ", clean_text)
            clean_text = re.sub(r"\s+", " ", clean_text).strip()
            result["text"] = clean_text
            result["pages"] = [{"section": "HTML Body", "text": clean_text}]

        elif extension in [".json", ".jsonl"]:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                raw_json = f.read()
            try:
                import json
                parsed = json.loads(raw_json)
                pretty_json = json.dumps(parsed, indent=2)
                result["text"] = pretty_json
                result["pages"] = [{"section": "JSON Data", "text": pretty_json}]
            except Exception:
                result["text"] = raw_json
                result["pages"] = [{"section": "Raw JSON", "text": raw_json}]

        elif extension in [".csv", ".tsv"]:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                csv_text = f.read()
            result["text"] = csv_text
            result["pages"] = [{"section": "Table Data", "text": csv_text}]

        elif extension in [".txt", ".md", ".markdown", ".rst", ".log", ".py", ".js", ".ts", ".jsx", ".tsx", ".yaml", ".yml", ".ini", ".env", ".c", ".cpp", ".h", ".sh", ".bat", ".ps1", ".sql"]:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            result["text"] = text
            result["pages"] = [{"section": "Full Document", "text": text}]

        else:
            logger.info(f"Reading {extension} with text reader.")
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            result["text"] = text
            result["pages"] = [{"section": "Text Content", "text": text}]

    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        result["text"] = ""
        result["metadata"]["error"] = str(e)

    return result